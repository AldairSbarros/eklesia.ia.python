from docx import Document

doc = Document()
doc.add_heading('A Trindade na Teologia Cristã', 0)
doc.add_paragraph('A doutrina da Trindade afirma que Deus é um só, em três pessoas: Pai, Filho e Espírito Santo.')
doc.add_paragraph('Esta verdade é fundamental para a fé cristã e está presente em todo o Novo Testamento.')
doc.save('j:/Projetos Python/eklesia.ia.python/tests/files/teste.docx')
