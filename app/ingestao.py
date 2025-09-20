# ingestao.py
from __future__ import annotations

import os
import re
import sys
import hashlib
import logging
import mimetypes
from datetime import datetime
from glob import glob
from typing import Optional, List

import fitz  # PyMuPDF
import docx  # python-docx
from bs4 import BeautifulSoup

from sqlalchemy import (
    create_engine, String, Text, DateTime, func, UniqueConstraint, Index
)
from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column, sessionmaker, Session

from dotenv import load_dotenv

# ----------------------------
# Configuração de logging
# ----------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger("ingestao")

# ----------------------------
# Ambiente / DB
# ----------------------------
load_dotenv()

DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql+psycopg2://aldai:2025@localhost:5432/eklesia_ia_db"
)
DOCS_DIR = os.getenv("DOCS_DIR", "./docs")

engine = create_engine(
    DATABASE_URL,
    pool_pre_ping=True,
    # Ajuste o pool se precisar de maior concorrência:
    # pool_size=5, max_overflow=10
)

SessionLocal = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)


class Base(DeclarativeBase):
    pass


class TextoBiblico(Base):
    __tablename__ = "textos_biblicos"

    id: Mapped[int] = mapped_column(primary_key=True, autoincrement=True)
    titulo: Mapped[str] = mapped_column(String(512), nullable=False)
    conteudo: Mapped[str] = mapped_column(Text, nullable=False)

    # Metadados úteis
    source_path: Mapped[Optional[str]] = mapped_column(String(1024))
    mime_type: Mapped[Optional[str]] = mapped_column(String(128))
    content_hash: Mapped[str] = mapped_column(String(64), nullable=False)

    created_at: Mapped[datetime] = mapped_column(
        DateTime(timezone=True), server_default=func.now()
    )

    __table_args__ = (
        UniqueConstraint("content_hash", name="uq_textos_biblicos_content_hash"),
        Index("ix_textos_biblicos_titulo", "titulo"),
    )


Base.metadata.create_all(engine)

# ----------------------------
# Utils
# ----------------------------
def normalize_text(text: str) -> str:
    """Normaliza quebras de linha e espaços, remove espaços em excesso nas bordas."""
    if not text:
        return ""
    # Normaliza \r\n -> \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove espaços à direita de cada linha
    text = "\n".join([ln.rstrip() for ln in text.split("\n")])
    # Evita mais de 3 quebras seguidas
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Tira espaços extras no começo e fim
    return text.strip()


def compute_hash(content: str, source_path: Optional[str]) -> str:
    """Gera um hash estável do conteúdo (e opcionalmente do caminho)."""
    h = hashlib.sha256()
    h.update((content or "").encode("utf-8", errors="ignore"))
    if source_path:
        h.update(source_path.encode("utf-8", errors="ignore"))
    return h.hexdigest()


def guess_mime(path: str) -> Optional[str]:
    mime, _ = mimetypes.guess_type(path)
    return mime


def detect_title_from_filename(path: str) -> str:
    base = os.path.basename(path)
    title = os.path.splitext(base)[0]
    # Substitui _ e - por espaço, capitaliza básico
    title = re.sub(r"[_\-]+", " ", title).strip()
    return title or "Documento"


# ----------------------------
# Extratores
# ----------------------------
def extrair_pdf(caminho: str) -> str:
    """Extrai texto de PDF usando PyMuPDF e faz fallback para pypdf se necessário."""
    # 1) Tenta com PyMuPDF
    try:
        with fitz.open(caminho) as doc:
            pages = []
            for page in doc:
                pages.append(page.get_text("text") or "")
            text = "\n\n".join(pages)
            text = normalize_text(text)
            if text:
                return text
    except Exception as e:
        logger.warning(f"[PDF] PyMuPDF falhou em '{caminho}': {e}")

    # 2) Fallback com pypdf (se instalado)
    try:
        from pypdf import PdfReader  # type: ignore
        with open(caminho, "rb") as f:
            reader = PdfReader(f)
            pages = []
            for p in reader.pages:
                pages.append((p.extract_text() or ""))
            text = "\n\n".join(pages)
            return normalize_text(text)
    except Exception as e:
        logger.error(f"[PDF] Fallback pypdf também falhou em '{caminho}': {e}")

    return ""


def extrair_docx(caminho: str) -> str:
    """Extrai texto de DOCX incluindo parágrafos e tabelas."""
    try:
        d = docx.Document(caminho)
    except Exception as e:
        logger.error(f"[DOCX] Falha abrindo '{caminho}': {e}")
        return ""

    parts: List[str] = []

    # Parágrafos
    for p in d.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    # Tabelas
    for table in d.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)

    return normalize_text("\n".join(parts))


def extrair_html(caminho: str) -> str:
    """Extrai texto limpo de HTML (remove scripts/estilos)."""
    try:
        with open(caminho, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()
    except Exception as e:
        logger.error(f"[HTML] Falha lendo '{caminho}': {e}")
        return ""

    try:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        # Remove linhas vazias em excesso
        text = re.sub(r"\n{3,}", "\n\n", text)
        return normalize_text(text)
    except Exception as e:
        logger.error(f"[HTML] Falha parseando '{caminho}': {e}")
        return ""


def extrair_txt(caminho: str) -> str:
    try:
        with open(caminho, "r", encoding="utf-8", errors="ignore") as f:
            return normalize_text(f.read())
    except Exception as e:
        logger.error(f"[TXT] Falha lendo '{caminho}': {e}")
        return ""


# ----------------------------
# Persistência
# ----------------------------
def salvar_texto(
    titulo: str,
    conteudo: str,
    source_path: Optional[str] = None,
    mime_type: Optional[str] = None,
) -> Optional[int]:
    conteudo = normalize_text(conteudo)
    if not conteudo:
        logger.info(f"[SKIP] Conteúdo vazio para '{source_path or titulo}'.")
        return None

    content_hash = compute_hash(conteudo, source_path)

    with SessionLocal() as session:  # type: Session
        # Evita duplicatas (idempotência)
        existing = session.query(TextoBiblico.id).filter(
            TextoBiblico.content_hash == content_hash
        ).first()
        if existing:
            logger.info(f"[DUPLICADO] Já existe registro com hash {content_hash[:10]}... (id={existing[0]})")
            return existing[0]

        obj = TextoBiblico(
            titulo=titulo[:512],
            conteudo=conteudo,
            source_path=source_path[:1024] if source_path else None,
            mime_type=mime_type[:128] if mime_type else None,
            content_hash=content_hash,
        )
        session.add(obj)
        session.commit()
        session.refresh(obj)
        logger.info(f"[OK] Salvo id={obj.id} '{titulo}'")
        return obj.id


# ----------------------------
# Pipeline
# ----------------------------
SUPPORTED_EXTS = {".pdf", ".docx", ".html", ".htm", ".txt", ".md"}


def processar_arquivo(path: str) -> Optional[int]:
    ext = os.path.splitext(path)[1].lower()
    mime = guess_mime(path)
    titulo = detect_title_from_filename(path)

    if ext == ".pdf":
        texto = extrair_pdf(path)
    elif ext == ".docx":
        texto = extrair_docx(path)
    elif ext in {".html", ".htm"}:
        texto = extrair_html(path)
    elif ext in {".txt", ".md"}:
        texto = extrair_txt(path)
    else:
        logger.info(f"[IGNORADO] Extensão não suportada: {ext} - {path}")
        return None

    if not texto:
        logger.warning(f"[VAZIO] Nada extraído de '{path}'")
        return None

    return salvar_texto(titulo=titulo, conteudo=texto, source_path=path, mime_type=mime)


def coletar_arquivos(root: str) -> List[str]:
    paths: List[str] = []
    patterns = ["**/*.pdf", "**/*.docx", "**/*.html", "**/*.htm", "**/*.txt", "**/*.md"]
    for pat in patterns:
        paths.extend(glob(os.path.join(root, pat), recursive=True))
    return paths


def main():
    # Permite passar um arquivo ou diretório na linha de comando
    target = sys.argv[1] if len(sys.argv) > 1 else DOCS_DIR

    if os.path.isdir(target):
        files = coletar_arquivos(target)
        if not files:
            logger.warning(f"Nenhum arquivo suportado encontrado em: {target}")
            return
        logger.info(f"Encontrados {len(files)} arquivos em {target}. Iniciando ingestão...")
        for i, path in enumerate(files, start=1):
            try:
                logger.info(f"[{i}/{len(files)}] Processando: {path}")
                processar_arquivo(path)
            except Exception as e:
                logger.exception(f"Falha ao processar '{path}': {e}")
    elif os.path.isfile(target):
        logger.info(f"Processando arquivo único: {target}")
        processar_arquivo(target)
    else:
        logger.error(f"Caminho não encontrado: {target}")


if __name__ == "__main__":
    main()
