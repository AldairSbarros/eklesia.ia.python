from docx import Document


def gerar_docx(sermao):
    doc = Document()
    doc.add_heading(sermao["tema"], level=1)
    for i, topico in enumerate(sermao["esboco"], 1):
        doc.add_heading(f"Tópico {i}", level=2)
        doc.add_paragraph(topico)
    doc.save("sermao.docx")
